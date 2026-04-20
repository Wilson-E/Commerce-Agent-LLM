"""Generate the ShopAssist AI final report for DS 440 Capstone."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(DOCS, "ShopAssist_AI_Final_Report.docx")

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


# ---------- Helpers (same pattern as generate_results_doc.py) ----------

def add_page_break():
    doc.add_page_break()


def center_title_page(title, subtitle, authors, advisor, course):
    for _ in range(8):
        doc.add_paragraph()
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("The Pennsylvania State University")
    r0.font.size = Pt(14)
    r0.font.name = "Calibri"
    r0.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p0b = doc.add_paragraph()
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0b = p0b.add_run("The College of Information Sciences and Technology")
    r0b.font.size = Pt(13)
    r0b.font.name = "Calibri"
    r0b.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p0c = doc.add_paragraph()
    p0c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0c = p0c.add_run(course)
    r0c.font.size = Pt(13)
    r0c.font.name = "Calibri"
    r0c.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(32)
    run.font.name = "Cambria"

    p1b = doc.add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1b = p1b.add_run(subtitle)
    r1b.font.size = Pt(16)
    r1b.font.name = "Calibri"
    r1b.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Prepared by")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(authors)
    r3.font.size = Pt(14)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Advisor: " + advisor)
    r4.font.size = Pt(12)
    r4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def centered_section_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(26)
    r.font.name = "Cambria"
    doc.add_paragraph()


def subheading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Cambria"
    r.bold = False


def sub_subheading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "Cambria"
    r.bold = True


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(11)


def body_no_indent(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(11)


def toc_entry(label, page, bold=False):
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(label)
    r.font.size = Pt(11)
    if bold:
        r.bold = True
    r2 = p.add_run("\t" + str(page))
    r2.font.size = Pt(11)
    if bold:
        r2.bold = True


def add_figure(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(caption)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def try_add_image(filename, width=Inches(6)):
    path = os.path.join(DOCS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
    else:
        body_no_indent(f"[Image not found: {filename}]")


# ======================================================================
# TITLE PAGE
# ======================================================================
center_title_page(
    "ShopAssist AI",
    "An AI-Native E-Commerce Shopping Assistant",
    "Bill Erdely, Arjun Shokeen",
    "Dr. Kaamran Raahemifar, PhD, PEng",
    "DS 440 Senior Capstone Project, Spring 2026",
)

# ======================================================================
# ABSTRACT
# ======================================================================
add_page_break()
centered_section_title("ABSTRACT")
body(
    "This report describes the background, development, and outcome of the ShopAssist AI "
    "web application created for a DS 440 Senior Capstone project at The Pennsylvania State "
    "University. The growing fragmentation of online shopping, where consumers must navigate "
    "dozens of tabs and rely on keyword searches that miss intent, motivates the need for a "
    "conversational alternative. The team developed a chat-first shopping platform that combines "
    "GPT-4 with a ReAct (Reasoning plus Acting) orchestration engine, Retrieval-Augmented "
    "Generation via a numpy-based vector store, and multi-source product aggregation from seven "
    "external APIs. The system classifies user intent, searches across sources using Reciprocal "
    "Rank Fusion, extracts user preferences for re-ranking, and enforces business-rule "
    "guardrails to prevent hallucinations, PII leakage, and prompt injection. The frontend is "
    "built with React, Vite, and Tailwind CSS, featuring WebSocket streaming, interactive "
    "product cards, a persistent cart grouped by merchant, and Stripe sandbox checkout. Security "
    "hardening includes JWT authentication with algorithm pinning, per-IP rate limiting, a "
    "token denylist, and first-frame WebSocket auth. Six crash test scenarios and fifteen "
    "automated security regression checks validate the platform's resilience. Several "
    "constraints, including time, free-tier API quotas, and team size, have left modifications "
    "for future work, including real merchant integrations, multi-modal input, mobile "
    "optimization, and international support."
)

# ======================================================================
# TABLE OF CONTENTS
# ======================================================================
add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TABLE OF CONTENTS")
r.font.size = Pt(22)
r.font.name = "Cambria"
doc.add_paragraph()

toc_entry("ABSTRACT", 2, bold=True)
toc_entry("TABLE OF CONTENTS", 3, bold=True)
toc_entry("LIST OF FIGURES", 4, bold=True)
toc_entry("INTRODUCTION", 5, bold=True)
toc_entry("    Problem Statement", 5)
toc_entry("    Motivation", 5)
toc_entry("    Objective", 6)
toc_entry("    End User Needs", 6)
toc_entry("LITERATURE SURVEY", 7, bold=True)
toc_entry("    Literature Review", 7)
toc_entry("    Assessment of available solutions and techniques", 8)
toc_entry("    Pros and Cons", 9)
toc_entry("REQUIREMENT SPECIFICATIONS", 10, bold=True)
toc_entry("    Market Requirement Analysis", 10)
toc_entry("    Design Requirement Analysis", 10)
toc_entry("    Constraints", 11)
toc_entry("    Assumptions", 12)
toc_entry("    Outcome Criteria", 12)
toc_entry("    Risks", 13)
toc_entry("SYSTEM DEVELOPMENT", 14, bold=True)
toc_entry("    Concept Generation", 14)
toc_entry("    System Planning", 15)
toc_entry("    Principle Operation", 16)
toc_entry("    Design Process", 17)
toc_entry("    System Design", 18)
toc_entry("    Functional Decomposition", 20)
toc_entry("TESTING", 25, bold=True)
toc_entry("    Attack 1: External API Failure", 25)
toc_entry("    Attack 2: Concurrent Users (DoS)", 26)
toc_entry("    Attack 3: Malicious Input", 27)
toc_entry("    Extended Test 4: Network Interruption", 28)
toc_entry("    Extended Test 5: Browser Compatibility", 29)
toc_entry("    Extended Test 6: Auth Churn", 29)
toc_entry("    Automated Security Regression Suite", 30)
toc_entry("CONCLUSION", 31, bold=True)
toc_entry("    Changes in development", 31)
toc_entry("    Challenges", 31)
toc_entry("    Lessons learned", 32)
toc_entry("FUTURE WORK", 33, bold=True)
toc_entry("    Real merchant integrations", 33)
toc_entry("    Multi-modal input (image and voice)", 33)
toc_entry("    Mobile-friendly application", 33)
toc_entry("    Optimization and improvement of UI elements", 34)
toc_entry("    Real-world testing and recommendation tuning", 34)
toc_entry("    International support", 34)
toc_entry("REFERENCES", 35, bold=True)

# ======================================================================
# LIST OF FIGURES
# ======================================================================
add_page_break()
centered_section_title("LIST OF FIGURES")
toc_entry("Figure 1: System Architecture", 18)
toc_entry("Figure 2: ReAct Orchestration Flow", 19)
toc_entry("Figure 3: Hybrid Search Pipeline", 20)
toc_entry("Figure 4: Category-Based API Routing", 21)
toc_entry("Figure 5: Functional Decomposition Diagram", 22)
toc_entry("Figure 6: PERT Chart", 15)
toc_entry("Figure 7: Critical Path Analysis", 16)
toc_entry("Figure 8: Testing Summary Table", 30)

# ======================================================================
# INTRODUCTION
# ======================================================================
add_page_break()
centered_section_title("INTRODUCTION")

subheading("Problem Statement")
body(
    "Online shopping has become the dominant mode of consumer commerce, yet the experience "
    "remains fragmented and frustrating. A consumer searching for a specific product, such as "
    "comfortable running shoes under eighty dollars, must open multiple browser tabs across "
    "Amazon, Google Shopping, specialty retailers, and comparison sites. Each site has its own "
    "search interface, its own filters, and its own ranking logic. The consumer is left to "
    "manually cross-reference results, compare prices, and track which products they have "
    "already evaluated. Traditional e-commerce search relies on keyword matching, which fails "
    "to capture the semantic intent behind a query. Typing \"comfortable wedding shoes\" returns "
    "results for each keyword separately, rather than understanding the combined concept of "
    "formal shoes with cushioning. There is no continuity across sessions or even across "
    "queries within the same session: every search starts from scratch, even if the user has "
    "already stated their preferred brand, budget, and size."
)
body(
    "The problem is compounded by the rise of large language models (LLMs), which have "
    "demonstrated remarkable capabilities in natural language understanding and generation. "
    "Consumers now expect conversational interfaces in other domains (customer support, "
    "productivity tools, coding assistants), but e-commerce has lagged behind. Existing "
    "AI-assisted shopping tools, such as Amazon Rufus and Shopify Sidekick, are locked to "
    "single retailers and cannot aggregate products across sources. There is a clear gap for "
    "a multi-source, conversational shopping assistant that understands intent, learns "
    "preferences, and enforces safety guardrails."
)

subheading("Motivation")
body(
    "The team's motivation for this project stemmed from a shared frustration with the "
    "current state of online shopping and a desire to apply recent advances in AI to a "
    "practical, user-facing application. As data science students, we had studied retrieval "
    "systems, natural language processing, and machine learning, but had not yet built a "
    "system that combined all three into a production-quality product. The ReAct framework "
    "(Yao et al., 2023) and Retrieval-Augmented Generation (Lewis et al., 2020) offered "
    "compelling architectures for building an AI agent that could reason about user queries, "
    "call external tools, and ground its responses in real product data."
)
body(
    "We were also motivated by the security and safety challenges inherent in LLM-based "
    "applications. Unlike a traditional web application, a system that generates natural "
    "language responses from user input is susceptible to prompt injection, hallucination, "
    "PII leakage, and other failure modes that standard web security measures do not address. "
    "We saw an opportunity to build not just a functional product but also a case study in "
    "how to harden an LLM-backed system for adversarial users."
)

subheading("Objective")
body(
    "The specific objectives of this project are to design and implement a chat-first "
    "e-commerce shopping assistant that: (1) accepts natural language queries and classifies "
    "user intent, (2) searches for products across multiple external APIs and a local vector "
    "store using hybrid semantic and keyword retrieval, (3) presents results as interactive "
    "product cards with size, color, and cart management controls, (4) learns and applies "
    "user preferences (brand, budget, size) within a session, (5) resolves conversational "
    "references such as \"it\" or \"the first one\" to specific product IDs, (6) enforces "
    "business-rule guardrails to prevent hallucinations, PII leakage, competitor mentions, "
    "and prompt injection, (7) provides a persistent cart grouped by merchant with Stripe "
    "sandbox checkout, and (8) supports multiple authentication methods including guest "
    "sessions, email/password, Google OAuth, and Microsoft OAuth."
)

subheading("End User Needs")
body(
    "The targeted end users are online shoppers who want a faster, more intuitive way to "
    "find products across multiple retailers. Specifically, the application addresses users "
    "who are frustrated by keyword-based search and want to describe what they need in "
    "natural language; users who shop across multiple sites and want a single interface that "
    "aggregates results; users who expect the system to remember their preferences within a "
    "session (for example, \"I prefer Nike and my budget is under $100\"); and users who are "
    "concerned about privacy and security when interacting with AI systems. The application "
    "is designed to be accessible to non-technical users through a clean chat interface, "
    "while providing enough depth (product attributes, ratings, stock status) to support "
    "informed purchase decisions."
)

# ======================================================================
# LITERATURE SURVEY
# ======================================================================
add_page_break()
centered_section_title("LITERATURE SURVEY")

subheading("Literature Review")
body(
    "The landscape of AI-assisted commerce has evolved rapidly since the introduction of "
    "transformer-based language models. At the core of our system is the ReAct framework "
    "(Yao et al., 2023), which interleaves reasoning traces with action steps in a language "
    "model. Unlike chain-of-thought prompting, which generates reasoning but takes no action, "
    "ReAct enables the model to call external tools (search APIs, databases, calculators) "
    "and observe their outputs before generating a final response. This pattern is particularly "
    "well-suited to e-commerce, where answering a user query often requires multiple steps: "
    "interpreting intent, searching for products, filtering results, and composing a natural "
    "language summary."
)
body(
    "Retrieval-Augmented Generation (Lewis et al., 2020) is the second foundational "
    "technique. RAG addresses a key limitation of LLMs: their training data is static and "
    "potentially outdated. By retrieving relevant documents (in our case, product descriptions) "
    "from an external knowledge base and injecting them into the model's context, RAG ensures "
    "that responses are grounded in current, factual data rather than relying solely on the "
    "model's parametric knowledge. Our implementation uses a numpy-based vector store with "
    "OpenAI embeddings and cosine similarity for retrieval, combined with keyword search "
    "across live APIs."
)
body(
    "The concept of conversational commerce, where consumers interact with an AI agent to "
    "discover and purchase products, has been explored by several industry players. Amazon's "
    "Rufus (launched 2024) is a shopping-focused chatbot integrated into the Amazon mobile "
    "app, capable of answering product questions and making recommendations. However, Rufus "
    "is limited to Amazon's own catalog and cannot aggregate results from external retailers. "
    "Shopify's Sidekick provides AI-powered assistance for merchants, not shoppers. Google's "
    "Shopping Graph and AI-enhanced product search aggregate listings from multiple retailers "
    "but present results in a traditional search-results format rather than a conversational "
    "interface. None of these solutions combine multi-source aggregation with conversational "
    "AI, user preference learning, and security guardrails in a single platform."
)
body(
    "Vector search and embedding-based retrieval have become standard tools for semantic "
    "information retrieval. Production systems typically use dedicated vector databases such "
    "as FAISS (Johnson et al., 2019), Milvus, or Pinecone, which offer optimized indexing "
    "for high-dimensional vectors. However, these systems introduce infrastructure complexity. "
    "For our project, we use a lightweight numpy-based approach that stores embeddings in "
    "memory and computes cosine similarity directly, eliminating the need for Docker or a "
    "standalone database service. This trade-off sacrifices scalability (the approach is "
    "practical for catalogs under ten thousand items) in favor of zero-infrastructure "
    "deployment."
)
body(
    "LLM guardrails and safety mechanisms have received increasing attention as language "
    "models are deployed in user-facing applications. NVIDIA's NeMo Guardrails (Rebedea et "
    "al., 2023) provides a programmable framework for controlling LLM behavior, including "
    "topic restriction, fact-checking, and output filtering. Guardrails AI offers a similar "
    "open-source framework. Our system implements custom guardrails tailored to the e-commerce "
    "domain: input-side filters for PII detection, competitor brand blocking, and off-topic "
    "rejection; and output-side filters for hallucinated discounts, unverified stock claims, "
    "and residual PII. The dual-layer approach ensures that even if one filter is bypassed "
    "(for example, through input paraphrasing), the output layer catches the violation."
)

subheading("Assessment of available solutions and techniques")
body(
    "Several approaches exist for building AI-powered shopping assistants, each with distinct "
    "trade-offs. The first category consists of single-retailer chatbots, such as Amazon Rufus "
    "and Walmart's AI assistant. These are tightly integrated with the retailer's catalog, "
    "search infrastructure, and purchase flow, providing a seamless experience within that "
    "ecosystem. However, they cannot help users compare products across retailers, which is "
    "one of the most common shopping tasks."
)
body(
    "The second category consists of multi-retailer aggregators, such as Google Shopping and "
    "PriceGrabber. These aggregate listings from many retailers and present them in a unified "
    "search interface. However, they use traditional keyword-based search and do not offer a "
    "conversational experience. Users cannot refine their search through dialogue, express "
    "complex preferences, or ask follow-up questions."
)
body(
    "The third category consists of general-purpose AI assistants, such as ChatGPT and "
    "Claude, which can answer shopping questions but lack real-time product data, cannot "
    "search external APIs, and cannot manage a cart or process payments. Their responses "
    "are based on training data rather than live inventory, making them unsuitable for "
    "actual purchase decisions."
)
body(
    "Our system, ShopAssist AI, occupies a fourth position: a multi-source conversational "
    "shopping assistant with real-time data, preference learning, and end-to-end purchase "
    "capability. It combines the conversational interface of general-purpose AI with the "
    "real-time data of aggregators and the purchase flow of single-retailer chatbots."
)

subheading("Pros and Cons")
body(
    "Single-retailer chatbots offer deep integration with one catalog, fast response times "
    "(no cross-network latency), and a seamless purchase flow. Their disadvantage is vendor "
    "lock-in: users cannot compare products across retailers, and the chatbot's "
    "recommendations are inherently biased toward the host retailer's inventory."
)
body(
    "Multi-retailer aggregators provide broad coverage and price comparison across many "
    "sources. Their disadvantage is the lack of conversational capability: users must "
    "express their needs through keywords and filters rather than natural language, and "
    "there is no session memory or preference learning."
)
body(
    "General-purpose AI assistants offer natural language understanding, multi-turn "
    "conversation, and broad knowledge. Their disadvantage is the absence of real-time "
    "product data, cart management, and payment processing. They are useful for research "
    "but not for completing a purchase."
)
body(
    "ShopAssist AI combines the strengths of all three categories: conversational interface "
    "with multi-turn memory (from general AI), multi-source aggregation with real-time data "
    "(from aggregators), and a complete purchase flow with cart and checkout (from "
    "single-retailer chatbots). Its disadvantage is the reliance on free-tier API quotas, "
    "which limit the volume of searches that can be performed in a given period, and the "
    "absence of real merchant partnerships for live inventory data."
)

# ======================================================================
# REQUIREMENT SPECIFICATIONS
# ======================================================================
add_page_break()
centered_section_title("REQUIREMENT SPECIFICATIONS")

subheading("Market Requirement Analysis")
body(
    "Global e-commerce sales surpassed $5.8 trillion in 2023 and are projected to reach "
    "$8 trillion by 2027 (eMarketer, 2024). Within this market, AI-powered shopping "
    "assistants represent a rapidly growing segment. A 2024 survey by Salesforce found that "
    "17% of consumers had already used a generative AI tool to assist with shopping "
    "decisions, and 45% expressed interest in doing so. The market for conversational "
    "commerce is estimated to grow at a compound annual growth rate of 23% through 2028 "
    "(Juniper Research, 2024)."
)
body(
    "The opportunity for ShopAssist AI lies in the gap between existing solutions. "
    "Single-retailer chatbots (Amazon Rufus) serve only their own catalog. Aggregators "
    "(Google Shopping) lack conversational capability. General-purpose AI (ChatGPT) lacks "
    "real-time product data. No existing product combines all three capabilities with "
    "security guardrails tailored to e-commerce. This gap is our target market."
)

subheading("Design Requirement Analysis")
body(
    "The design of ShopAssist AI must accommodate several core requirements. First, the "
    "system must support natural language product search, meaning users should be able to "
    "describe what they want in plain English and receive relevant results. This requires "
    "both semantic understanding (via embeddings) and real-time data access (via external "
    "APIs)."
)
body(
    "Second, the system must present products in an interactive format, not just text. "
    "Product cards with images, ratings, size/color selectors, and add-to-cart buttons are "
    "necessary for users to make informed decisions without leaving the chat interface."
)
body(
    "Third, the system must maintain conversation context, including the ability to resolve "
    "references (\"add it to my cart\"), remember stated preferences (\"I like Nike\"), and "
    "compress long conversations to stay within token limits."
)
body(
    "Fourth, the system must enforce security at multiple levels: authentication (JWT with "
    "algorithm pinning), input validation (PII redaction, off-topic blocking), output "
    "validation (hallucination prevention, competitor blocking), and transport security "
    "(CSP headers, CORS restrictions)."
)
body(
    "Fifth, the system must support a complete purchase flow: browsing, adding to cart, "
    "reviewing cart contents (grouped by merchant), and completing payment through a "
    "third-party processor (Stripe)."
)

subheading("Constraints")
body(
    "Several constraints shaped the development of ShopAssist AI. The most significant is "
    "budget. All external APIs used in the project operate on free tiers with limited quotas. "
    "The OpenAI API (GPT-4o-mini) provides limited daily tokens on the free tier; SerpAPI, "
    "RapidAPI, Rainforest, and ScraperAPI each impose their own rate limits. This constraint "
    "forced the team to implement category-based routing, which directs queries to only the "
    "most relevant APIs for a given product category, conserving quota for genuinely "
    "conversational interactions."
)
body(
    "The second constraint is time. The project was developed over a single semester (roughly "
    "fourteen weeks of active development). This limited the number of features that could "
    "be implemented and, in particular, compressed the security hardening phase, which alone "
    "consumed approximately two weeks."
)
body(
    "The third constraint is team size. With only two developers, the project required "
    "careful division of labor between frontend and backend work, with both team members "
    "contributing to integration testing and security review."
)
body(
    "The fourth constraint is the in-memory architecture. To simplify deployment and "
    "eliminate infrastructure dependencies (no Docker, no PostgreSQL, no Redis), the system "
    "stores all state in memory. This means that server restarts clear all user data, cart "
    "contents, and conversation history. While acceptable for a demo and capstone project, "
    "a production deployment would require a persistent database."
)

subheading("Assumptions")
body(
    "The project is developed under several assumptions. First, the team assumes that the "
    "OpenAI API and external product APIs will remain available and stable throughout the "
    "development and demonstration period. Second, the team assumes that users have a modern "
    "web browser with WebSocket support (Chrome, Firefox, Safari, Edge). Third, the team "
    "assumes that the free-tier API quotas are sufficient for demonstration purposes, though "
    "not for sustained production traffic. Fourth, the team assumes that the Stripe sandbox "
    "environment accurately simulates real payment processing for demonstration purposes."
)

subheading("Outcome Criteria")
body(
    "The criteria used to determine success of the project include: whether a natural "
    "language query entered in the chat interface returns relevant product results from "
    "multiple sources; whether the system correctly classifies user intent and routes to "
    "the appropriate handler (conversational response or product search); whether interactive "
    "product cards with size/color selection and add-to-cart functionality render correctly "
    "in the chat; whether the cart persists across queries and groups items by merchant; "
    "whether Stripe checkout creates a valid payment session; whether guardrails correctly "
    "block PII, competitor brands, off-topic queries, and prompt injection attempts; and "
    "whether the system remains operational after all six crash test scenarios."
)

subheading("Risks")
body(
    "The primary risk is API quota exhaustion during testing or demonstration. If the "
    "OpenAI free tier is exhausted, the system cannot classify intent, generate responses, "
    "or create embeddings. Mitigation: category-based routing reduces LLM calls, and the "
    "embedding cache eliminates the need for re-embedding on restart."
)
body(
    "The second risk is LLM hallucination. Despite guardrails, the model may occasionally "
    "generate incorrect product information (wrong price, fabricated features). Mitigation: "
    "the output guardrails engine scrubs responses for specific hallucination patterns, and "
    "product data is always sourced from the API or local catalog, not generated by the model."
)
body(
    "The third risk is prompt injection. A malicious user could attempt to extract API keys, "
    "the system prompt, or data from other users' sessions. Mitigation: the guardrails engine "
    "filters known injection patterns on input, the system prompt strictly constrains the "
    "model's role, and JWT-based identity isolation prevents cross-user data access."
)
body(
    "The fourth risk is WebSocket stability under poor network conditions. A dropped "
    "connection mid-stream could leave the UI in a frozen state. Mitigation: the frontend "
    "implements exponential backoff reconnection, handler nulling before socket close, "
    "and a debounced connection-lost banner."
)

# ======================================================================
# SYSTEM DEVELOPMENT
# ======================================================================
add_page_break()
centered_section_title("SYSTEM DEVELOPMENT")

subheading("Concept Generation")
body(
    "The team set out to build an application at the intersection of two rapidly advancing "
    "fields: large language models and e-commerce. The initial concept was simple: what if "
    "a shopper could describe what they want in natural language and receive product "
    "recommendations from multiple retailers, all within a single conversational interface? "
    "This concept was refined through several rounds of brainstorming and technical "
    "prototyping."
)
body(
    "Early prototyping explored whether existing chatbot frameworks (Rasa, Dialogflow) "
    "could support the level of reasoning required for multi-step shopping queries. These "
    "frameworks proved too rigid: they required predefined intents and entities, and could "
    "not dynamically decide when to call external tools. The ReAct framework (Yao et al., "
    "2023), which uses the LLM itself to decide which tools to call and when, offered a "
    "more flexible alternative."
)
body(
    "The team also considered whether to use a managed vector database (Milvus, Pinecone) "
    "or a lightweight local approach. Milvus was prototyped initially but required Docker "
    "and a standalone service, adding deployment complexity. The team pivoted to a numpy-based "
    "approach that stores embeddings in memory and computes cosine similarity directly. This "
    "eliminated an entire infrastructure layer while remaining performant for the project's "
    "catalog size (under 200 products)."
)

subheading("System Planning")
body(
    "The project followed an agile development methodology with approximately two-week "
    "sprints. The team used Git for version control, with both members committing to the "
    "same repository and resolving conflicts through pull requests. The development was "
    "divided into four main phases."
)
body(
    "Phase 1 (weeks 1 through 4) focused on the backend foundation: building the ReAct "
    "orchestration loop, integrating the OpenAI API for intent classification and response "
    "generation, setting up the numpy vector store with OpenAI embeddings, and connecting "
    "the first external product APIs (SerpAPI and RapidAPI)."
)
body(
    "Phase 2 (weeks 5 through 8) focused on the frontend: migrating from Create React App "
    "to Vite with Tailwind CSS, building the chat interface with WebSocket streaming, "
    "designing and implementing interactive product cards, and creating the cart drawer "
    "with merchant grouping."
)
body(
    "Phase 3 (weeks 9 through 11) focused on integration and polish: wiring the frontend "
    "to backend REST endpoints, adding Stripe sandbox checkout, implementing coreference "
    "resolution and preference extraction, and adding additional APIs (Rainforest, ScraperAPI, "
    "ASOS, Home Depot, Open Food Facts)."
)
body(
    "Phase 4 (weeks 12 through 14) focused on security hardening and testing: conducting an "
    "adversarial review, implementing JWT authentication with algorithm pinning, adding "
    "guardrails for input and output, rate limiting guest sessions, performing six crash test "
    "scenarios, and writing fifteen automated security regression checks."
)

try_add_image("PERT_Chart_v5.png", width=Inches(6.5))
add_figure("Figure 6: PERT Chart showing task dependencies and project timeline.")

try_add_image("Critical_Path_Analysis_v5.png", width=Inches(6.5))
add_figure("Figure 7: Critical Path Analysis identifying the longest sequential task chain.")

subheading("Principle Operation")
body(
    "The principle operation of ShopAssist AI begins when the user navigates to the web "
    "application in their browser. Upon loading, the frontend presents a login page with "
    "four options: guest access, email/password registration, Google OAuth, and Microsoft "
    "OAuth. Selecting any option generates a JWT token that authenticates all subsequent "
    "requests."
)
body(
    "After authentication, the user enters the chat interface. A WebSocket connection is "
    "established using a first-frame auth protocol: the client sends a JSON message "
    "containing the token type and JWT as the first frame, and the server validates it "
    "before accepting further messages."
)
body(
    "When the user types a message and presses send, the message is transmitted over the "
    "WebSocket to the backend's orchestration engine. The engine uses a two-call architecture. "
    "The first LLM call, which has no tools attached, classifies the message as conversational "
    "or shopping-related. If conversational (greeting, small talk, general question), the "
    "model responds directly. If shopping-related, the engine enters a ReAct loop."
)
body(
    "In the ReAct loop, the LLM has access to six tools: search_products, "
    "get_product_details, add_to_cart, get_cart, get_order_status, and browse_category. "
    "The model decides which tool to call, the executor runs the tool, and the result is "
    "fed back to the model for the next iteration. This loop continues for up to five "
    "iterations. At each iteration, the response is streamed token by token over the "
    "WebSocket, giving the user real-time feedback."
)
body(
    "When the search_products tool is invoked, it triggers the hybrid search pipeline. "
    "The query is simultaneously sent to the numpy vector store (semantic search via cosine "
    "similarity) and to the relevant external APIs (keyword search, routed by product "
    "category). Results from both sources are merged using Reciprocal Rank Fusion, then "
    "re-ranked based on any user preferences extracted from the conversation. The final "
    "product list is returned to the LLM, which formats it as a natural language response "
    "with embedded product data that the frontend renders as interactive cards."
)
body(
    "Users can interact with product cards directly (selecting size, color, adding to cart) "
    "or continue the conversation (\"tell me more about the first one\", \"add it to my "
    "cart\"). The system resolves these references through coreference resolution, which "
    "uses the LLM to map pronouns and ordinal references to specific product IDs."
)
body(
    "When the user is ready to purchase, they open the cart drawer, review items grouped "
    "by merchant, and click \"Proceed to Checkout.\" The frontend sends the cart to the "
    "backend, which creates a Stripe Checkout Session with line items, tax, and a "
    "success/cancel redirect URL. The user is redirected to Stripe's hosted payment page "
    "to complete the transaction."
)

subheading("Design Process")
body(
    "The design process followed an iterative, agile approach with several significant "
    "pivots. The most consequential design decisions are summarized below."
)
body(
    "Vector database: the initial design used Milvus, a production-grade vector database "
    "that requires Docker and a standalone service. After experiencing deployment complexity "
    "and startup latency, the team pivoted to a numpy-based vector store that computes "
    "cosine similarity in memory. Embeddings are generated via the OpenAI API on first "
    "startup and cached to a local JSON file for subsequent launches."
)
body(
    "Frontend framework: the initial frontend used Create React App (CRA). After CRA's "
    "slow dev server and large bundle sizes began costing iteration speed, the team migrated "
    "to Vite with Tailwind CSS. The migration took approximately one week and resulted in "
    "a significantly faster development loop."
)
body(
    "Payment processing: the original scope did not include payment. After implementing "
    "the cart, the team decided to add Stripe sandbox checkout to demonstrate a complete "
    "purchase flow. This required adding Stripe integration on both the backend (Checkout "
    "Session creation, webhook handling) and frontend (redirect to Stripe's hosted page)."
)
body(
    "Security: the original design did not include authentication. After an adversarial "
    "review midway through the project exposed that any client could impersonate any user "
    "by self-declaring a user ID, the team implemented a full JWT-based auth system with "
    "algorithm pinning, token denylist, per-IP rate limiting, and first-frame WebSocket "
    "auth. This pivot consumed approximately two weeks."
)

subheading("System Design")
body(
    "The system is organized into three layers: the Interaction Layer, the Reasoning Layer, "
    "and the Services/Data Layer."
)
body(
    "The Interaction Layer consists of the React frontend, the FastAPI REST API, and the "
    "WebSocket streaming endpoint. The frontend communicates with the backend through REST "
    "calls (for authentication, cart management, and product listing) and a persistent "
    "WebSocket connection (for chat streaming). The REST API provides endpoints for guest "
    "session creation, user registration, login (email, Google, Microsoft), cart operations "
    "(add, remove, update, merge), product listing, and Stripe checkout session creation."
)
body(
    "The Reasoning Layer contains the orchestration engine, which implements the ReAct "
    "pattern. It receives user messages from the WebSocket handler, passes them through "
    "input guardrails, classifies intent via a first LLM call (no tools), and either "
    "responds directly (conversational) or enters a tool-calling loop (shopping). The "
    "guardrails engine operates at both ends: input guardrails filter PII, competitor "
    "brands, and off-topic queries; output guardrails scrub hallucinated discounts, "
    "unverified stock claims, and residual PII."
)
body(
    "The Services/Data Layer provides the underlying data access. The vector database "
    "service manages numpy-based embeddings for semantic search. The product database "
    "service handles the local catalog and coordinates calls to seven external APIs, "
    "routing by product category. The memory service manages conversation history, "
    "coreference resolution, preference extraction, and context compression. The user "
    "database service stores in-memory cart and profile state. The auth database service "
    "stores credentials and manages JWT issuance and validation."
)

subheading("Functional Decomposition")
body(
    "The system's functionality is decomposed into eight major modules, each responsible "
    "for a distinct capability."
)

sub_subheading("Orchestration Engine (orchestrator.py)")
body(
    "The orchestration engine is the central coordinator of the system. It receives "
    "a user message and session context, then executes a two-call architecture. The first "
    "call sends the message to the LLM with no tools attached, using a router prompt that "
    "instructs the model to either respond conversationally or output the keyword "
    "SEARCH_PRODUCTS. If the model outputs SEARCH_PRODUCTS, the engine enters a ReAct loop "
    "with tools enabled. The loop iterates up to five times: at each step, the model either "
    "calls a tool (and the executor runs it) or generates a final text response. All "
    "responses are streamed token by token over the WebSocket."
)

sub_subheading("Guardrails Engine (guardrails.py)")
body(
    "The guardrails engine enforces business rules at both input and output. On input, "
    "it checks for competitor brand mentions (Amazon Basics, Kirkland, Target Brand, "
    "Walmart Brand) using keyword matching, detects off-topic queries (recipes, weather, "
    "jokes, medical advice) using regex patterns, and redacts PII (SSN patterns matching "
    "NNN-NN-NNNN, credit card numbers matching sixteen-digit patterns, and email addresses "
    "matching standard patterns) by replacing them with [REDACTED]. On output, it scrubs "
    "any PII that the LLM may have generated, removes fabricated discount percentages, "
    "strips competitor brand names, and flags unverified stock claims."
)

sub_subheading("Hybrid Search Pipeline (vector_db.py, executor.py)")
body(
    "The search pipeline operates in four stages. First, the user query is embedded via "
    "the OpenAI embeddings API and compared against the product embedding matrix using "
    "cosine similarity, returning the top N semantically similar products. Second, the "
    "query is sent to external APIs (routed by product category) for keyword-based results. "
    "Third, results from both sources are merged using Reciprocal Rank Fusion (RRF), which "
    "assigns each result a score of 1/(k + rank), where k = 60, and sums scores across "
    "sources. Fourth, if the user has stated preferences (brand, budget, size), matching "
    "products receive a preference boost that re-ranks them toward the top."
)

sub_subheading("Multi-Source API Aggregation (product_db.py, search_intent.py)")
body(
    "The product database service coordinates searches across seven external APIs: SerpAPI "
    "(Google Shopping), RapidAPI (real-time product search), Rainforest API (Amazon data), "
    "ScraperAPI (Amazon structured data), ASOS (fashion via RapidAPI), Home Depot (home "
    "improvement via RapidAPI), and Open Food Facts (food/grocery, free, no key required). "
    "To conserve free-tier quota, the search intent module classifies each query into a "
    "product category (technology, fashion, home, food, general) and routes it to only the "
    "relevant APIs. Results from all sources pass through a result aggregator that merges, "
    "deduplicates (using Jaccard similarity on product names), and ranks the unified list."
)

sub_subheading("Conversational Intelligence (memory.py)")
body(
    "The memory service manages three aspects of conversational intelligence. Coreference "
    "resolution handles references like \"it,\" \"that,\" \"the first one,\" and \"the Nike "
    "ones\" by sending the conversation history and the reference to the LLM, which returns "
    "the specific product ID being referenced. Preference extraction scans each user message "
    "for brand mentions, budget phrases (\"under $100,\" \"less than fifty dollars\"), and size "
    "signals (\"I wear size 10\"), storing extracted preferences per user for use in search "
    "re-ranking. Context compression monitors the total token count of the conversation "
    "history; when it exceeds 8,000 tokens, the service sends older messages to the LLM for "
    "summarization, replacing them with a condensed summary to stay within limits."
)

sub_subheading("Authentication and Security (auth.py, main.py)")
body(
    "The authentication system supports four methods: email/password registration with "
    "bcrypt hashing, Google OAuth, Microsoft OAuth, and anonymous guest sessions. All "
    "methods issue JWT tokens with the HS256 algorithm. The JWT validation layer pins the "
    "algorithm (rejecting alg:none and key-confusion attacks), requires four claims (exp, "
    "sub, jti, iat), and checks the jti against an in-memory denylist populated on logout. "
    "Guest tokens carry a 59-minute TTL and are rate-limited to 10 per IP per 60 seconds. "
    "When a guest registers, the /api/cart/merge endpoint transfers their cart to the new "
    "account. Security middleware adds CSP, X-Frame-Options, X-Content-Type-Options, and "
    "Referrer-Policy headers to every response. CORS is restricted to concrete origin URLs."
)

sub_subheading("Payment Processing (main.py checkout endpoints)")
body(
    "The checkout system integrates with Stripe in sandbox mode. When the user clicks "
    "\"Proceed to Checkout,\" the frontend sends the cart to the backend, which creates "
    "a Stripe Checkout Session with line items, a configurable tax rate (default 8%), and "
    "success/cancel redirect URLs. The user is redirected to Stripe's hosted payment page. "
    "Upon successful payment, Stripe sends a webhook to the backend, which creates an order "
    "record per merchant and clears the user's cart. Webhook events include idempotency "
    "handling to prevent duplicate order creation."
)

sub_subheading("Frontend (React, Vite, Tailwind, Zustand)")
body(
    "The frontend is built with React 18 and Vite 5.0 for fast builds, Tailwind CSS 4.0 "
    "for utility-first styling with dark/light mode support, and Zustand for state "
    "management. Three Zustand stores manage auth state (JWT token, user profile), cart "
    "state (items with variants, persisted to localStorage), and theme state (dark/light "
    "mode preference). The chat interface connects via WebSocket and streams responses "
    "token by token. Product cards render inline with chat messages and include size/color "
    "selectors, ratings, stock status, and add-to-cart buttons. The cart drawer slides out "
    "from the right, groups items by merchant, and provides quantity controls. The checkout "
    "page shows an order summary with subtotals, tax, and total before redirecting to Stripe."
)

# ======================================================================
# TESTING
# ======================================================================
add_page_break()
centered_section_title("TESTING")

body(
    "The testing strategy for ShopAssist AI encompasses six crash test scenarios designed "
    "to probe the system's resilience under adversarial conditions, plus a suite of fifteen "
    "automated security regression checks. The crash tests target three failure categories: "
    "platform dependency (external API failures), scalability (concurrent user load), and "
    "security (malicious input and prompt injection). Three extended tests cover network "
    "interruption, browser compatibility, and rapid authentication churn."
)

subheading("Attack 1: External API Failure and Cascade Timeout")
body(
    "This test validates that the platform handles external API failures gracefully. The "
    "system depends on multiple third-party APIs (SerpAPI, RapidAPI, Rainforest, ScraperAPI, "
    "ASOS, Home Depot) for live product data. If one or more of these APIs go down, become "
    "slow, or return errors, the platform must continue operating rather than crashing."
)
body(
    "The test procedure involves sending a baseline product search query with the internet "
    "connected, then disconnecting the internet and sending another query. The platform "
    "should either return results from the local product catalog (fallback) or display a "
    "friendly error message. After reconnecting, a third query confirms recovery."
)
body(
    "Defensive measures include category-aware API routing (only calling APIs relevant to "
    "the product category), parallel API calls with a 15-second timeout per source, a local "
    "product fallback catalog (sample_products.json), and graceful exception handling that "
    "logs errors without crashing the server process."
)

subheading("Attack 2: Rapid Concurrent Users (DoS Simulation)")
body(
    "This test validates that the platform handles concurrent load without crashing or "
    "becoming unresponsive. The procedure involves opening six browser tabs simultaneously, "
    "each pointed at the application, and rapidly sending different product search queries "
    "across all tabs."
)
body(
    "Defensive measures include FastAPI's async architecture (which handles many concurrent "
    "WebSocket connections without blocking), a ReAct loop iteration cap of five per query "
    "(preventing any single request from consuming the server indefinitely), a context token "
    "limit of 8,000 tokens (preventing unbounded memory growth), and session isolation "
    "(each WebSocket connection has its own state)."
)

subheading("Attack 3: Malicious and Oversized User Input")
body(
    "This test encompasses six sub-tests: (A) oversized input (a message with the phrase "
    "\"Find me a laptop\" repeated 50+ times), (B) prompt injection attempting to extract "
    "API keys, (C) prompt injection attempting to extract the system prompt, (D) PII "
    "injection (SSN, credit card, email), (E) competitor brand queries (Amazon Basics, "
    "Kirkland), and (F) off-topic requests (recipe for chocolate cake). A post-attack "
    "verification query confirms that the platform remains fully operational."
)
body(
    "Defensive measures include the GuardrailsEngine's input checking (competitor brand "
    "blocking, off-topic detection, PII redaction), output checking (residual PII scrubbing, "
    "hallucination prevention), a constrained system prompt that strictly defines the "
    "assistant's role, the ReAct loop iteration cap, and the context token limit."
)

subheading("Extended Test 4: Network Interruption Mid-Stream")
body(
    "This test validates that the platform handles a network dropout during a streaming "
    "response without freezing the UI, duplicating replies on reconnect, or leaking "
    "half-open backend connections. The procedure involves sending a query, toggling Wi-Fi "
    "off as the response streams, waiting eight seconds, then toggling Wi-Fi back on."
)
body(
    "Defensive measures include nulling all WebSocket event handlers on the old socket "
    "before opening a new one (preventing a late-firing onopen from clobbering the new "
    "socket's heartbeat state), removing the auto-retry of pending messages (the server "
    "may have already processed the request), and debouncing the connection-lost banner "
    "by 1.5 seconds to avoid flicker on fast reconnects."
)

subheading("Extended Test 5: Browser and Platform Matrix")
body(
    "This test confirms that the platform is not silently Chrome-only. The procedure "
    "involves performing a guest login and chat query on Chrome, Safari, and Firefox, "
    "verifying that the landing page renders correctly, the WebSocket connects, and the "
    "chat response appears within five seconds on all three browsers."
)
body(
    "Defensive measures include environment-driven URLs (no localhost hardcoded in frontend "
    "source), CORS restricted to a concrete origin list, and sessionStorage access wrapped "
    "in try/catch to handle Safari private-mode quota limits."
)

subheading("Extended Test 6: Rapid Authentication Churn")
body(
    "This test validates that logging out and back in repeatedly (seven cycles of guest "
    "access) does not leak cart contents, chat history, or pending messages between "
    "identities, and does not leak WebSocket connections on the backend."
)
body(
    "Defensive measures include server-minted guest UUIDs (two guests never collide), "
    "JWT-derived user_id on all endpoints (switching tokens switches identity server-side), "
    "handler nulling before WebSocket close (preventing stale sockets from invoking state "
    "updates), and no auto-retry of pending messages on reconnect."
)

subheading("Automated Security Regression Suite")
body(
    "In addition to the manual crash tests, the project includes a suite of fifteen "
    "automated security regression checks implemented in test_phase1_security.py. These "
    "checks cover: JWT algorithm pinning (rejecting alg:none tokens), wrong-key JWT "
    "rejection, cross-user cart isolation (one user cannot read another's cart), logout "
    "replay protection (a revoked token cannot be reused), invalid session_id format "
    "rejection, oversized message rejection, tampered checkout price detection, and "
    "guest rate-limit enforcement. The suite runs as a standard Python test file and is "
    "executed before each deployment."
)

# ======================================================================
# CONCLUSION (reused verbatim from generate_results_doc.py)
# ======================================================================
add_page_break()
centered_section_title("CONCLUSION")

subheading("Changes in development")
body(
    "Our development process focused on the same end goal throughout the semester, but we made "
    "major changes in the way specific tasks were accomplished. The original plan used Milvus as "
    "the vector database, which required Docker and a standalone service, along with Redis for "
    "conversation memory. Both were shelved in favor of a lightweight numpy based local vector "
    "store and an in memory cache, which eliminated an entire layer of infrastructure and "
    "simplified deployment considerably. Another significant shift happened midway through the "
    "project when we migrated the frontend from Create React App to Vite with Tailwind, after "
    "CRA's slow dev server and large bundle sizes started costing us iteration speed. On the "
    "backend, we also added Stripe sandbox checkout after initial scoping focused only on chat "
    "and cart features."
)
body(
    "The most substantial change came late in the project, after an adversarial security review "
    "exposed gaps in our authentication and WebSocket layer. We introduced JWT based auth with "
    "short lived guest tokens, per IP rate limiting on guest issuance, a token denylist for "
    "logout, and a first frame WebSocket auth protocol that moved tokens out of query strings. "
    "We also added category based routing on the backend to conserve free tier OpenAI quota, "
    "which kept the demo functional even under frequent traffic during testing sessions."
)

subheading("Challenges")
body(
    "A consistent challenge across the project was the free tier quota on the OpenAI API. Each "
    "chat request consumed embedding and completion tokens, and during heavy testing we "
    "repeatedly hit the daily cap. This forced us to design a category based routing layer that "
    "intercepts obvious queries (for example, plain category listings) and answers them without "
    "an LLM call, saving quota for genuinely conversational turns. If budget had not been a "
    "constraint, we could have used the more capable GPT 4 tier throughout and stress tested "
    "the system with automated load generators."
)
body(
    "Another major challenge was learning the WebSocket resiliency patterns needed for a "
    "streaming chat interface. Early versions suffered from subtle bugs: late firing onopen "
    "handlers clobbered a fresh socket's heartbeat reference, an auto retry loop could double "
    "bill users if the server had already processed a message, and the connection lost banner "
    "flickered on every brief reconnect. Fixing these required carefully nulling event handlers "
    "before closing sockets, removing the auto retry path in favor of a user facing timeout, "
    "and debouncing the banner by roughly a second and a half."
)
body(
    "Time was the other persistent challenge. We underestimated how much of the schedule would "
    "be spent on security hardening, specifically the full pass after the adversarial review, "
    "which touched JWT validation, WebSocket auth, cart merge flows, content security policy "
    "headers, and input validation on every schema. While we delivered the feature set we "
    "scoped, the security pass alone consumed roughly two weeks we had originally planned to "
    "spend on UI polish and additional merchant categories."
)

subheading("Lessons learned")
body(
    "One of the first things we learned was the value of operating in an agile loop rather than "
    "front loading the plan. We made architectural changes (Milvus out, Stripe in, CRA to Vite) "
    "at points in the semester where a rigid plan would have blocked us. Git was central to "
    "this: we could work on the code simultaneously, revert bad directions quickly, and keep a "
    "clean history of why each shift happened."
)
body(
    "A second lesson was the importance of an explicit security pass, separate from feature "
    "work. Before our adversarial review, the backend accepted any self declared user id on the "
    "chat endpoint, which meant a malicious client could impersonate anyone and read their "
    "cart. Writing a regression test suite specifically for security (JWT algorithm pinning, "
    "cross user cart isolation, logout replay protection, oversized payload rejection) caught "
    "several classes of bug that functional tests would never have surfaced."
)
body(
    "A third lesson was that LLM backed systems need guardrails sized to the threat model, not "
    "just to the happy path. Early prompt injection tests showed that a sufficiently crafty "
    "user could coax the model into mentioning competitor brands, fabricating discounts, or "
    "exposing PII from other sessions. We added both input filters (category and off topic "
    "blocking) and output scrubbing (PII, promo codes, competitor names), and we learned that "
    "the output layer is the one that actually protects users, since inputs can be paraphrased "
    "endlessly."
)

# ======================================================================
# FUTURE WORK (reused verbatim from generate_results_doc.py)
# ======================================================================
add_page_break()
centered_section_title("FUTURE WORK")

body(
    "Modifications for the Commerce Agent LLM have been left for future work due to time "
    "constraints. Throughout our semester, we built the system from scratch and learned many "
    "things along the way. Our main focus was to deliver a chat first shopping experience with "
    "a working cart, Stripe sandboxed checkout, and hardened authentication, which we "
    "accomplished. However, with more time, the application has meaningful room for expansion "
    "in several directions."
)

subheading("Real merchant integrations")
body(
    "The product catalog in the current build is a curated JSON file of roughly forty items "
    "distributed across six synthetic merchant brands. A production deployment would integrate "
    "directly with real e commerce platforms such as Shopify, Magento, and BigCommerce, and "
    "ingest live inventory, pricing, and variant data. This would require building per platform "
    "adapters, a normalization layer for their differing schemas, and a webhook pipeline to "
    "keep our vector index fresh as merchants update their catalogs."
)

subheading("Multi-modal input (image and voice)")
body(
    "The chat interface is text only today. A natural extension is image search (for example, "
    "find shoes that look like this photo) using CLIP style embeddings alongside the existing "
    "text embeddings, and voice input via the Web Speech API or an OpenAI Whisper endpoint. "
    "Both would require changes to the orchestrator so that a single user turn can include "
    "multiple modalities, and to the vector database so that image and text embeddings live in "
    "comparable spaces."
)

subheading("Mobile-friendly application")
body(
    "We built the assistant to work well on desktop browsers. The current UI scales to mobile "
    "web but is not optimized for it: the product grid crowds on narrow screens and the cart "
    "drawer overlaps the keyboard on iOS. A dedicated mobile web layout, and eventually a React "
    "Native build for iOS and Android, would significantly broaden the addressable user base. "
    "The React Native port could reuse the Zustand stores and the REST client with only the "
    "component tree rewritten."
)

subheading("Optimization and improvement of UI elements")
body(
    "The team worked to make the interface presentable, but several polish items were cut for "
    "time. Specifically, we would add skeleton loaders on the product grid while the WebSocket "
    "is streaming, a persistent assistant is thinking indicator during the ReAct loop, and "
    "optimistic UI on the cart drawer so that add to cart feels instant even when the backend "
    "call takes two hundred milliseconds. We would also replace the placeholder category icons "
    "with real product images once a hosting solution for user uploaded assets is in place."
)

subheading("Real-world testing and recommendation tuning")
body(
    "One of the biggest limitations of our platform at time of writing is that the "
    "recommendation engine has been exercised only by the development team and a small pool of "
    "testers. Production grade relevance tuning needs real user behavior: click through rates "
    "on suggested products, add to cart conversion per query type, and search abandonment "
    "signals. We would instrument the frontend with privacy respecting event logging, feed "
    "those signals back into a reranking layer on top of the current hybrid retrieval, and A/B "
    "test prompt variations against matched user cohorts."
)

subheading("International support")
body(
    "The application is currently United States only: prices are in dollars, tax is a flat "
    "eight percent stand in for U.S. sales tax, and all product copy and assistant responses "
    "are in English. An international launch would require multi currency display and "
    "conversion, per region tax calculation (VAT, GST, provincial combinations), and "
    "localization of both the UI chrome and the assistant's responses. The assistant's "
    "localization is the interesting part: it would need either translated system prompts per "
    "language or a single multilingual model capable of maintaining persona across languages."
)

# ======================================================================
# REFERENCES
# ======================================================================
add_page_break()
centered_section_title("REFERENCES")

refs = [
    "[1] S. Yao, J. Zhao, D. Yu, N. Du, I. Shafran, K. Narasimhan, and Y. Cao, "
    "\"ReAct: Synergizing Reasoning and Acting in Language Models,\" in Proc. ICLR 2023, "
    "2023.",

    "[2] P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, H. Kuttler, "
    "M. Lewis, W. Yih, T. Rocktaschel, S. Riedel, and D. Kiela, \"Retrieval-Augmented "
    "Generation for Knowledge-Intensive NLP Tasks,\" in Proc. NeurIPS 2020, 2020.",

    "[3] J. Johnson, M. Douze, and H. Jegou, \"Billion-scale similarity search with GPUs,\" "
    "IEEE Trans. Big Data, vol. 7, no. 3, pp. 535-547, 2021.",

    "[4] E. Rebedea, C. Dinu, N. Sreedhar, C. Parisien, and J. Cohen, \"NeMo Guardrails: "
    "A Toolkit for Controllable and Safe LLM Applications with Programmable Rails,\" "
    "arXiv preprint arXiv:2310.10501, 2023.",

    "[5] OpenAI, \"GPT-4 Technical Report,\" arXiv preprint arXiv:2303.08774, 2023.",

    "[6] OpenAI, \"OpenAI API Documentation,\" 2024. [Online]. Available: "
    "https://platform.openai.com/docs",

    "[7] S. Ramaswamy, \"FastAPI Documentation,\" 2024. [Online]. Available: "
    "https://fastapi.tiangolo.com",

    "[8] Meta Platforms, \"React Documentation,\" 2024. [Online]. Available: "
    "https://react.dev",

    "[9] Stripe, \"Stripe API Documentation,\" 2024. [Online]. Available: "
    "https://stripe.com/docs/api",

    "[10] SerpAPI, \"Google Shopping API Documentation,\" 2024. [Online]. Available: "
    "https://serpapi.com/google-shopping-api",

    "[11] eMarketer, \"Global Ecommerce Forecast 2024,\" Insider Intelligence, 2024.",

    "[12] Juniper Research, \"Conversational Commerce: Key Trends and Market Forecasts "
    "2024-2028,\" Juniper Research Ltd., 2024.",

    "[13] Salesforce, \"State of the Connected Customer, 6th Edition,\" Salesforce "
    "Research, 2024.",

    "[14] A. Radford, J. Wu, R. Child, D. Luan, D. Amodei, and I. Sutskever, \"Language "
    "Models are Unsupervised Multitask Learners,\" OpenAI, 2019.",

    "[15] T. Brown, B. Mann, N. Ryder, M. Subbiah, J. Kaplan, P. Dhariwal, et al., "
    "\"Language Models are Few-Shot Learners,\" in Proc. NeurIPS 2020, 2020.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    r.font.size = Pt(10)
    r.font.name = "Calibri"


# ======================================================================
# Save
# ======================================================================
doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
