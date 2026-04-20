"""Generate Results, Conclusions, Future Work Word doc for Commerce-Agent-LLM."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUTPUT = "/Users/billyerdely/Documents/Penn_State/Commerce-Agent-LLM/docs/Results_Conclusions_Future_Work.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def add_page_break():
    doc.add_page_break()

def center_title_page(title, authors):
    for _ in range(12):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(32)
    run.font.name = "Cambria"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(authors)
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def centered_section_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(26)
    r.font.name = "Cambria"
    doc.add_paragraph()

def subheading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.name = "Cambria"
    r.bold = False

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(11)

def toc_entry(label, page, bold=False):
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(label)
    r.font.size = Pt(11)
    if bold:
        r.bold = True
    r2 = p.add_run("\t" + str(page))
    r2.font.size = Pt(11)
    if bold:
        r2.bold = True

center_title_page("Results, Conclusions, Future Work", "Bill Erdely, Arjun Shokeen")

add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TABLE OF CONTENTS")
r.font.size = Pt(22)
r.font.name = "Cambria"
doc.add_paragraph()

toc_entry("TABLE OF CONTENTS", 2, bold=True)
toc_entry("CONCLUSIONS", 3, bold=True)
toc_entry("    Changes in development", 3)
toc_entry("    Challenges", 3)
toc_entry("    Lessons learned", 4)
toc_entry("FUTURE WORK", 5, bold=True)
toc_entry("    Real merchant integrations", 5)
toc_entry("    Multi-modal input (image and voice)", 5)
toc_entry("    Mobile-friendly application", 5)
toc_entry("    Optimization and improvement of UI elements", 6)
toc_entry("    Real-world testing and recommendation tuning", 6)
toc_entry("    International support", 6)

add_page_break()
centered_section_title("CONCLUSIONS")

subheading("Changes in development")
body(
    "Our development process focused on the same end goal throughout the semester, but we made "
    "major changes in the way specific tasks were accomplished. The original plan used Milvus as "
    "the vector database, which required Docker and a standalone service, along with Redis for "
    "conversation memory. Both were shelved in favor of a lightweight numpy based local vector "
    "store and an in memory cache, which eliminated an entire layer of infrastructure and "
    "simplified deployment considerably. Another significant shift happened midway through the "
    "project when we migrated the frontend from Create React App to Vite with Tailwind, after "
    "CRA's slow dev server and large bundle sizes started costing us iteration speed. On the "
    "backend, we also added Stripe sandbox checkout after initial scoping focused only on chat "
    "and cart features."
)
body(
    "The most substantial change came late in the project, after an adversarial security review "
    "exposed gaps in our authentication and WebSocket layer. We introduced JWT based auth with "
    "short lived guest tokens, per IP rate limiting on guest issuance, a token denylist for "
    "logout, and a first frame WebSocket auth protocol that moved tokens out of query strings. "
    "We also added category based routing on the backend to conserve free tier OpenAI quota, "
    "which kept the demo functional even under frequent traffic during testing sessions."
)

subheading("Challenges")
body(
    "A consistent challenge across the project was the free tier quota on the OpenAI API. Each "
    "chat request consumed embedding and completion tokens, and during heavy testing we "
    "repeatedly hit the daily cap. This forced us to design a category based routing layer that "
    "intercepts obvious queries (for example, plain category listings) and answers them without "
    "an LLM call, saving quota for genuinely conversational turns. If budget had not been a "
    "constraint, we could have used the more capable GPT 4 tier throughout and stress tested "
    "the system with automated load generators."
)
body(
    "Another major challenge was learning the WebSocket resiliency patterns needed for a "
    "streaming chat interface. Early versions suffered from subtle bugs: late firing onopen "
    "handlers clobbered a fresh socket's heartbeat reference, an auto retry loop could double "
    "bill users if the server had already processed a message, and the connection lost banner "
    "flickered on every brief reconnect. Fixing these required carefully nulling event handlers "
    "before closing sockets, removing the auto retry path in favor of a user facing timeout, "
    "and debouncing the banner by roughly a second and a half."
)
body(
    "Time was the other persistent challenge. We underestimated how much of the schedule would "
    "be spent on security hardening, specifically the full pass after the adversarial review, "
    "which touched JWT validation, WebSocket auth, cart merge flows, content security policy "
    "headers, and input validation on every schema. While we delivered the feature set we "
    "scoped, the security pass alone consumed roughly two weeks we had originally planned to "
    "spend on UI polish and additional merchant categories."
)

subheading("Lessons learned")
body(
    "One of the first things we learned was the value of operating in an agile loop rather than "
    "front loading the plan. We made architectural changes (Milvus out, Stripe in, CRA to Vite) "
    "at points in the semester where a rigid plan would have blocked us. Git was central to "
    "this: we could work on the code simultaneously, revert bad directions quickly, and keep a "
    "clean history of why each shift happened."
)
body(
    "A second lesson was the importance of an explicit security pass, separate from feature "
    "work. Before our adversarial review, the backend accepted any self declared user id on the "
    "chat endpoint, which meant a malicious client could impersonate anyone and read their "
    "cart. Writing a regression test suite specifically for security (JWT algorithm pinning, "
    "cross user cart isolation, logout replay protection, oversized payload rejection) caught "
    "several classes of bug that functional tests would never have surfaced."
)
body(
    "A third lesson was that LLM backed systems need guardrails sized to the threat model, not "
    "just to the happy path. Early prompt injection tests showed that a sufficiently crafty "
    "user could coax the model into mentioning competitor brands, fabricating discounts, or "
    "exposing PII from other sessions. We added both input filters (category and off topic "
    "blocking) and output scrubbing (PII, promo codes, competitor names), and we learned that "
    "the output layer is the one that actually protects users, since inputs can be paraphrased "
    "endlessly."
)

add_page_break()
centered_section_title("FUTURE WORK")

body(
    "Modifications for the Commerce Agent LLM have been left for future work due to time "
    "constraints. Throughout our semester, we built the system from scratch and learned many "
    "things along the way. Our main focus was to deliver a chat first shopping experience with "
    "a working cart, Stripe sandboxed checkout, and hardened authentication, which we "
    "accomplished. However, with more time, the application has meaningful room for expansion "
    "in several directions."
)

subheading("Real merchant integrations")
body(
    "The product catalog in the current build is a curated JSON file of roughly forty items "
    "distributed across six synthetic merchant brands. A production deployment would integrate "
    "directly with real e commerce platforms such as Shopify, Magento, and BigCommerce, and "
    "ingest live inventory, pricing, and variant data. This would require building per platform "
    "adapters, a normalization layer for their differing schemas, and a webhook pipeline to "
    "keep our vector index fresh as merchants update their catalogs."
)

subheading("Multi-modal input (image and voice)")
body(
    "The chat interface is text only today. A natural extension is image search (for example, "
    "find shoes that look like this photo) using CLIP style embeddings alongside the existing "
    "text embeddings, and voice input via the Web Speech API or an OpenAI Whisper endpoint. "
    "Both would require changes to the orchestrator so that a single user turn can include "
    "multiple modalities, and to the vector database so that image and text embeddings live in "
    "comparable spaces."
)

subheading("Mobile-friendly application")
body(
    "We built the assistant to work well on desktop browsers. The current UI scales to mobile "
    "web but is not optimized for it: the product grid crowds on narrow screens and the cart "
    "drawer overlaps the keyboard on iOS. A dedicated mobile web layout, and eventually a React "
    "Native build for iOS and Android, would significantly broaden the addressable user base. "
    "The React Native port could reuse the Zustand stores and the REST client with only the "
    "component tree rewritten."
)

subheading("Optimization and improvement of UI elements")
body(
    "The team worked to make the interface presentable, but several polish items were cut for "
    "time. Specifically, we would add skeleton loaders on the product grid while the WebSocket "
    "is streaming, a persistent assistant is thinking indicator during the ReAct loop, and "
    "optimistic UI on the cart drawer so that add to cart feels instant even when the backend "
    "call takes two hundred milliseconds. We would also replace the placeholder category icons "
    "with real product images once a hosting solution for user uploaded assets is in place."
)

subheading("Real-world testing and recommendation tuning")
body(
    "One of the biggest limitations of our platform at time of writing is that the "
    "recommendation engine has been exercised only by the development team and a small pool of "
    "testers. Production grade relevance tuning needs real user behavior: click through rates "
    "on suggested products, add to cart conversion per query type, and search abandonment "
    "signals. We would instrument the frontend with privacy respecting event logging, feed "
    "those signals back into a reranking layer on top of the current hybrid retrieval, and A/B "
    "test prompt variations against matched user cohorts."
)

subheading("International support")
body(
    "The application is currently United States only: prices are in dollars, tax is a flat "
    "eight percent stand in for U.S. sales tax, and all product copy and assistant responses "
    "are in English. An international launch would require multi currency display and "
    "conversion, per region tax calculation (VAT, GST, provincial combinations), and "
    "localization of both the UI chrome and the assistant's responses. The assistant's "
    "localization is the interesting part: it would need either translated system prompts per "
    "language or a single multilingual model capable of maintaining persona across languages."
)

doc.save(OUTPUT)
print(f"wrote {OUTPUT}")
